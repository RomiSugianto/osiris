import io

from django.contrib.auth.decorators import permission_required, login_required
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.forms import inlineformset_factory
from django.http import HttpResponse, HttpResponseNotFound
from django.shortcuts import render, get_object_or_404, redirect
from django.utils.decorators import method_decorator
from django.views.generic import ListView
from docx import Document
from docx.shared import Inches
from docxtpl import DocxTemplate

from inventory.forms import *
from inventory.models import *


# Create your views here.

@login_required()
def index(request):
    return render(request, 'inventory/index.html')


@login_required()
def sp_view(request, pk):
    data = get_object_or_404(RequestServer, pk=pk)
    doc = DocxTemplate("inventory/templates/inventory/doc/template-sp.docx")
    doc.render({'data': data})

    doc_io = io.BytesIO()  # create a file-like object
    doc.save(doc_io)  # save data to file-like object
    doc_io.seek(0)  # go to the beginning of the file-like object

    response = HttpResponse(doc_io.read())

    # Content-Disposition header makes a file downloadable
    response["Content-Disposition"] = "attachment; filename=Surat-Perjanjian-'%s'.docx" % data

    # Set the appropriate Content-Type for docx file
    response["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return response


@login_required()
def fp_view(request, pk):
    data = get_object_or_404(RequestServer, pk=pk)
    doc = DocxTemplate("inventory/templates/inventory/doc/template-fp.docx")
    doc.render({'data': data})

    doc_io = io.BytesIO()  # create a file-like object
    doc.save(doc_io)  # save data to file-like object
    doc_io.seek(0)  # go to the beginning of the file-like object

    docx = Document(doc_io)
    tables = docx.tables
    p = tables[1].rows[1].cells[0].add_paragraph()
    r = p.add_run()
    r.add_picture(data.created_by.profile.signature, width=Inches(1.0), height=Inches(1.0))
    # docx.save(doc_io)
    docx_io = io.BytesIO()
    docx.save(docx_io)
    docx_io.seek(0)

    response = HttpResponse(docx_io.read())

    # Content-Disposition header makes a file downloadable
    response["Content-Disposition"] = "attachment; filename=Form-Permohonan-'%s'.docx" % data

    # Set the appropriate Content-Type for docx file
    response["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return response


def pie_chart(request):
    labels = []
    data = []

    total = RequestServer.objects.count()
    queryset = User.objects.all()
    for user in queryset:
        label = user.get_full_name()
        count = user.requestserver_set.count()
        labels.append(label)
        data.append(count)

    return render(request, 'pie_chart.html', {
        'labels': labels,
        'data': data,
        'total': total,
    })


def test_view(request, pk):
    request_server = get_object_or_404(RequestServer, id=pk)
    # return HttpResponse(request_server)
    form = RequestServerForm(request.POST or None, instance=request_server)
    if form.is_valid():
        form.save()
        return redirect('RequestServerList')
    return render(request, 'inventory/page_test.html', {'form': form})


class RequestServerView(ListView):
    template_name = 'inventory/request_server.html'
    model = RequestServer
    paginate_by = 10
    ordering = ['-created_at']

    @method_decorator(permission_required('inventory.view_requestserver'))
    def dispatch(self, *args, **kwargs):
        return super().dispatch(*args, **kwargs)


def request_server_approved(request):
    request_server_list = RequestServer.objects.filter(approved_by__isnull=False, approved_at__isnull=False).order_by(
        '-created_at')
    pc_list = Pc.objects.filter(requestserver__isnull=True)
    return render(request, 'inventory/request_server_approved.html',
                  {'object_list': request_server_list, 'pc_list': pc_list})


def request_server_set_pc(request, pk):
    request_server = get_object_or_404(RequestServer, pk=pk)
    form = RequestServerSetPcForm(request.POST or None, instance=request_server)
    if form.is_valid():
        instance = form.save(commit=False)
        instance.pc_status = 'ready'
        instance.save()
        return redirect('RequestServerApprovedList')


@permission_required('inventory.add_requestserver')
def request_server_new(request):
    if request.method == 'GET':
        form = RequestServerNewForm(initial={'created_by': get_object_or_404(User, pk=request.user.id)})
        return render(request, 'inventory/request_server_new.html', {'form': form})
    elif request.method == 'POST':
        form = RequestServerNewForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('RequestServerList')


@permission_required('inventory.approve_requestserver', raise_exception=True)
def request_server_approve(request, pk):
    if request.method == 'POST':
        request_server = get_object_or_404(RequestServer, pk=pk)
        if request.POST['remark'] is not None:
            request_server.remark = request.POST['remark']
            request_server.approved_by = get_object_or_404(User, pk=request.POST['user_id'])
            request_server.approved_at = timezone.now()
            request_server.pc_status = 'prepare'
            request_server.save()
            return redirect('RequestServerList')


def request_server_receive_document(request, pk):
    request_server = get_object_or_404(RequestServer, pk=pk)
    request_server.document_received_back_at = timezone.now()
    request_server.pc_status = 'completed'
    request_server.save()
    return redirect('RequestServerList')


def request_server_send(request, pk):
    request_server = get_object_or_404(RequestServer, pk=pk)
    request_server.document_sent_at = timezone.now()
    request_server.sent_at = timezone.now()
    request_server.pc_status = 'sent'
    request_server.save()
    return redirect('RequestServerApprovedList')


class CustomerListView(ListView):
    template_name = 'inventory/customer.html'
    model = Customer
    paginate_by = 10
    ordering = ['id']

    @method_decorator(permission_required('inventory.view_customer'))
    def dispatch(self, *args, **kwargs):
        return super().dispatch(*args, **kwargs)


@permission_required('inventory.add_customer', raise_exception=True)
def customer_new(request):
    if request.method == 'POST':
        form = CustomerForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('CustomerList')
    else:
        form = CustomerForm()
        return render(request, 'inventory/customer_new.html', {'form': form})


@permission_required('inventory.change_customer', raise_exception=True)
def customer_edit(request, pk):
    customer = get_object_or_404(Customer, pk=pk)
    form = CustomerForm(request.POST or None, instance=customer)
    if form.is_valid():
        form.save()
        return redirect('CustomerList')
    return render(request, 'inventory/customer_edit.html', {'form': form})


@permission_required('inventory.delete_customer', raise_exception=True)
def customer_delete(request, pk):
    customer = get_object_or_404(Customer, pk=pk)
    customer.delete()
    return redirect('CustomerList')


class PrincipalListView(ListView):
    template_name = 'inventory/principal.html'
    model = Principal
    paginate_by = 10
    ordering = ['id']

    @method_decorator(permission_required('inventory.view_principal'))
    def dispatch(self, *args, **kwargs):
        return super().dispatch(*args, **kwargs)


@permission_required('inventory.add_principal', raise_exception=True)
def principal_new(request):
    if request.method == 'POST':
        form = PrincipalForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('PrincipalList')
    else:
        form = PrincipalForm()
        return render(request, 'inventory/principal_new.html', {'form': form})


@permission_required('inventory.change_principal', raise_exception=True)
def principal_edit(request, pk):
    principal = get_object_or_404(Principal, pk=pk)
    form = PrincipalForm(request.POST or None, instance=principal)
    if form.is_valid():
        form.save()
        return redirect('PrincipalList')
    return render(request, 'inventory/principal_edit.html', {'form': form})


@permission_required('inventory.delete_principal', raise_exception=True)
def principal_delete(request, pk):
    principal = get_object_or_404(Principal, pk=pk)
    principal.delete()
    return redirect('PrincipalList')


class DistributorListView(ListView):
    template_name = 'inventory/distributor.html'
    model = Distributor
    paginate_by = 10
    ordering = ['id']

    @method_decorator(permission_required('inventory.view_distributor'))
    def dispatch(self, *args, **kwargs):
        return super().dispatch(*args, **kwargs)


@permission_required('inventory.add_distributor', raise_exception=True)
def distributor_new(request):
    if request.method == 'POST':
        form = DistributorForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('DistributorList')
    else:
        form = DistributorForm()
        return render(request, 'inventory/distributor_new.html', {'form': form})


@permission_required('inventory.change_distributor', raise_exception=True)
def distributor_edit(request, pk):
    distributor = get_object_or_404(Distributor, pk=pk)
    form = DistributorForm(request.POST or None, instance=distributor)
    if form.is_valid():
        form.save()
        return redirect('DistributorList')
    return render(request, 'inventory/distributor_edit.html', {'form': form})


@permission_required('inventory.delete_distributor', raise_exception=True)
def distributor_delete(request, pk):
    distributor = get_object_or_404(Distributor, pk=pk)
    distributor.delete()
    return redirect('DistributorList')


class AntiVirusView(ListView):
    template_name = 'inventory/antivirus.html'
    model = AntiVirus
    paginate_by = 10
    ordering = '-id'


def anti_virus_new(request):
    if request.method == 'POST':
        form = AntiVirusForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('AntiVirusList')
    else:
        form = AntiVirusForm
        return render(request, 'inventory/antivirus_new.html', {'form': form})


def anti_virus_delete(request, pk):
    anti_virus = get_object_or_404(AntiVirus, pk=pk)
    anti_virus.delete()
    return redirect('AntiVirusList')


def anti_virus_edit(request, pk):
    anti_virus = get_object_or_404(AntiVirus, pk=pk)
    form = AntiVirusForm(request.POST or None, instance=anti_virus)
    if form.is_valid():
        form.save()
        return redirect('AntiVirusList')
    return render(request, 'inventory/antivirus_edit.html', {'form': form})


class PcView(ListView):
    template_name = 'inventory/pc.html'
    model = Pc
    paginate_by = 10
    ordering = '-pc_number'


def pc_delete(request, pk):
    pc = get_object_or_404(Pc, pk=pk)
    pc.delete()
    return redirect('PcList')


def pc_edit(request, pk):
    pc = get_object_or_404(Pc, pk=pk)
    form = PcForm(request.POST or None, instance=pc)
    if form.is_valid():
        form.save()
        return redirect('PcList')
    return render(request, 'inventory/pc_edit.html', {'form': form})


def pc_new(request):
    if request.method == 'POST':
        form = PcForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('PcList')
    else:
        form = PcForm
        return render(request, 'inventory/pc_new.html', {'form': form})


class SsdView(ListView):
    template_name = 'inventory/ssd.html'
    model = Ssd
    paginate_by = 10


def ssd_delete(request, pk):
    ssd = get_object_or_404(Ssd, pk=pk)
    ssd.delete()
    return redirect('SsdList')


def ssd_edit(request, pk):
    ssd = get_object_or_404(Ssd, pk=pk)
    form = SsdForm(request.POST or None, instance=ssd)
    if form.is_valid():
        form.save()
        return redirect('SsdList')
    return render(request, 'inventory/ssd_edit.html', {'form': form})


def ssd_new(request):
    if request.method == 'POST':
        form = SsdForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('SsdList')
    else:
        form = SsdForm
        return render(request, 'inventory/ssd_new.html', {'form': form})


class HddView(ListView):
    template_name = 'inventory/hdd.html'
    model = Hdd
    paginate_by = 10


def hdd_delete(request, pk):
    ssd = get_object_or_404(Ssd, pk=pk)
    ssd.delete()
    return redirect('SsdList')


def hdd_edit(request, pk):
    ssd = get_object_or_404(Hdd, pk=pk)
    form = HddForm(request.POST or None, instance=ssd)
    if form.is_valid():
        form.save()
        return redirect('HddList')
    return render(request, 'inventory/hdd_edit.html', {'form': form})


def hdd_new(request):
    if request.method == 'POST':
        form = HddForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('HddList')
    else:
        form = HddForm
        return render(request, 'inventory/hdd_new.html', {'form': form})


class PsuView(ListView):
    template_name = 'inventory/psu.html'
    model = Psu
    paginate_by = 10


def psu_delete(request, pk):
    psu = get_object_or_404(Psu, pk=pk)
    psu.delete()
    return redirect('PsuList')


def psu_edit(request, pk):
    psu = get_object_or_404(Psu, pk=pk)
    form = PsuForm(request.POST or None, instance=psu)
    if form.is_valid():
        form.save()
        return redirect('PsuList')
    return render(request, 'inventory/psu_edit.html', {'form': form})


def psu_new(request):
    if request.method == 'POST':
        form = PsuForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('PsuList')
    else:
        form = PsuForm
        return render(request, 'inventory/psu_new.html', {'form': form})


class RamView(ListView):
    template_name = 'inventory/ram.html'
    model = Ram
    paginate_by = 10


def ram_delete(request, pk):
    ram = get_object_or_404(Ram, pk=pk)
    ram.delete()
    return redirect('RamList')


def ram_edit(request, pk):
    ram = get_object_or_404(Ram, pk=pk)
    form = RamForm(request.POST or None, instance=ram)
    if form.is_valid():
        form.save()
        return redirect('RamList')
    return render(request, 'inventory/ram_edit.html', {'form': form})


def ram_new(request):
    if request.method == 'POST':
        form = RamForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('RamList')
    else:
        form = RamForm
        return render(request, 'inventory/ram_new.html', {'form': form})


class MotherboardView(ListView):
    template_name = 'inventory/motherboard.html'
    model = Motherboard
    paginate_by = 10


def motherboard_delete(request, pk):
    motherboard = get_object_or_404(Motherboard, pk=pk)
    motherboard.delete()
    return redirect('MotherboardList')


def motherboard_edit(request, pk):
    motherboard = get_object_or_404(Motherboard, pk=pk)
    form = MotherboardForm(request.POST or None, instance=motherboard)
    if form.is_valid():
        form.save()
        return redirect('MotherboardList')
    return render(request, 'inventory/motherboard_edit.html', {'form': form})


def motherboard_new(request):
    if request.method == 'POST':
        form = MotherboardForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('MotherboardList')
    else:
        form = MotherboardForm
        return render(request, 'inventory/motherboard_new.html', {'form': form})